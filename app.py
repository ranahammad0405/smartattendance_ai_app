import streamlit as st
import os
import json
from datetime import datetime
from groq import Groq
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re

# ──────────────────────────────────────────────
# PAGE CONFIG
# ──────────────────────────────────────────────
st.set_page_config(
    page_title="SmartNotes - AI Study Notes Generator",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ──────────────────────────────────────────────
# CUSTOM CSS — Enhanced Light Mode, Polished UI
# ──────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Serif+Display:ital@0;1&family=DM+Sans:wght@300;400;500;600;700&display=swap');

:root {
    --bg:        #f5f4f0;
    --surface:   #ffffff;
    --card2:     #f0eeea;
    --border:    #e2ddd8;
    --accent:    #2563eb;
    --accent2:   #7c3aed;
    --text:      #1a1814;
    --muted:     #6b6560;
    --success:   #059669;
    --warning:   #d97706;
    --shadow:    0 1px 3px rgba(0,0,0,0.08), 0 4px 16px rgba(0,0,0,0.06);
    --shadow-lg: 0 8px 32px rgba(0,0,0,0.12);
}

html, body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > section,
[data-testid="stMain"],
[data-testid="stMainBlockContainer"],
.main .block-container,
.stApp {
    background-color: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
}

*, *::before, *::after { box-sizing: border-box; }

p, span, div, li, td, th, label, small {
    color: var(--text);
    font-family: 'DM Sans', sans-serif;
}

h1, h2, h3, h4, h5, h6 {
    font-family: 'DM Serif Display', serif !important;
    color: var(--text) !important;
    letter-spacing: -0.02em !important;
}

/* SIDEBAR */
[data-testid="stSidebar"],
[data-testid="stSidebar"] > div,
[data-testid="stSidebarContent"] {
    background-color: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}

[data-testid="stSidebar"] p,
[data-testid="stSidebar"] span,
[data-testid="stSidebar"] div,
[data-testid="stSidebar"] label,
[data-testid="stSidebar"] small,
[data-testid="stSidebar"] h1,
[data-testid="stSidebar"] h2,
[data-testid="stSidebar"] h3,
[data-testid="stSidebar"] h4,
[data-testid="stSidebar"] .stMarkdown {
    color: var(--text) !important;
}

[data-testid="stSidebar"] input,
[data-testid="stSidebar"] textarea {
    background: #f8f8f6 !important;
    color: var(--text) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 8px !important;
}

[data-testid="stSidebar"] [data-testid="stMetricValue"],
[data-testid="stSidebar"] [data-testid="stMetricLabel"] {
    color: var(--text) !important;
}

/* TABS */
[data-testid="stTabs"] [data-baseweb="tab-list"] {
    background: var(--card2) !important;
    border-radius: 12px !important;
    padding: 4px !important;
    border: 1px solid var(--border) !important;
    gap: 2px !important;
}

[data-testid="stTabs"] [data-baseweb="tab"] {
    border-radius: 9px !important;
    padding: 0.5rem 1.2rem !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 500 !important;
    font-size: 0.9rem !important;
    color: #4b5563 !important;
    background: transparent !important;
    transition: all 0.15s ease !important;
}

[data-testid="stTabs"] [data-baseweb="tab"]:hover {
    color: var(--accent) !important;
    background: rgba(37,99,235,0.06) !important;
}

[data-testid="stTabs"] [aria-selected="true"] {
    background: var(--surface) !important;
    color: var(--accent) !important;
    box-shadow: var(--shadow) !important;
}

[data-testid="stTabs"] [data-baseweb="tab-panel"] {
    background: transparent !important;
    padding-top: 1rem !important;
}

/* BUTTONS */
.stButton > button {
    background: var(--accent) !important;
    color: #ffffff !important;
    border: none !important;
    border-radius: 10px !important;
    padding: 0.65rem 1.5rem !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.95rem !important;
    letter-spacing: 0.01em !important;
    transition: all 0.18s ease !important;
    box-shadow: 0 2px 8px rgba(37,99,235,0.28) !important;
}

.stButton > button:hover {
    background: #1d4ed8 !important;
    color: #ffffff !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(37,99,235,0.38) !important;
}

.stButton > button:active {
    transform: translateY(0px) !important;
    background: #1e40af !important;
    color: #ffffff !important;
}

.stButton > button span,
.stButton > button p { color: #ffffff !important; }

/* DOWNLOAD BUTTONS */
[data-testid="stDownloadButton"] > button {
    background: #ffffff !important;
    color: var(--accent) !important;
    border: 2px solid var(--accent) !important;
    border-radius: 10px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-weight: 600 !important;
    font-size: 0.93rem !important;
    transition: all 0.18s ease !important;
}

[data-testid="stDownloadButton"] > button span,
[data-testid="stDownloadButton"] > button p { color: var(--accent) !important; }

[data-testid="stDownloadButton"] > button:hover {
    background: var(--accent) !important;
    color: #ffffff !important;
    transform: translateY(-1px) !important;
    box-shadow: 0 6px 20px rgba(37,99,235,0.25) !important;
}

[data-testid="stDownloadButton"] > button:hover span,
[data-testid="stDownloadButton"] > button:hover p { color: #ffffff !important; }

/* INPUTS */
[data-testid="stTextInput"] input,
[data-testid="stTextArea"] textarea,
input[type="text"],
textarea {
    background: #ffffff !important;
    color: var(--text) !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 10px !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.93rem !important;
    caret-color: var(--accent) !important;
}

[data-testid="stTextInput"] input:focus,
[data-testid="stTextArea"] textarea:focus {
    border-color: var(--accent) !important;
    box-shadow: 0 0 0 3px rgba(37,99,235,0.12) !important;
    outline: none !important;
}

[data-testid="stTextInput"] label,
[data-testid="stTextArea"] label,
[data-testid="stSelectbox"] label,
.stTextInput label, .stTextArea label {
    color: var(--text) !important;
    font-weight: 600 !important;
    font-size: 0.88rem !important;
    font-family: 'DM Sans', sans-serif !important;
}

input::placeholder, textarea::placeholder {
    color: #9ca3af !important;
    opacity: 1 !important;
}

/* SELECT BOX — force white bg + dark text everywhere */
[data-testid="stSelectbox"] div[data-baseweb="select"],
[data-testid="stSelectbox"] div[data-baseweb="select"] > div,
[data-testid="stSelectbox"] div[data-baseweb="select"] > div > div,
[data-testid="stSelectbox"] [data-baseweb="select"] [role="combobox"],
[data-testid="stSelectbox"] [data-baseweb="select"] input {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #1a1814 !important;
    border: 1.5px solid var(--border) !important;
    border-radius: 10px !important;
}

/* The visible selected value text */
[data-testid="stSelectbox"] [data-baseweb="select"] [data-baseweb="select"] span,
[data-testid="stSelectbox"] [data-baseweb="select"] span,
[data-testid="stSelectbox"] [data-baseweb="select"] div[class*="ValueContainer"] span,
[data-testid="stSelectbox"] [data-baseweb="select"] div[class*="singleValue"],
[data-testid="stSelectbox"] [data-baseweb="select"] div[class*="placeholder"] {
    color: #1a1814 !important;
}

/* Dropdown menu list */
[data-baseweb="popover"] ul,
[data-baseweb="menu"],
[data-baseweb="menu"] ul,
[data-baseweb="menu"] li,
[data-baseweb="popover"] [role="option"],
[data-baseweb="popover"] [data-baseweb="menu-item"] {
    background: #ffffff !important;
    background-color: #ffffff !important;
    color: #1a1814 !important;
}

[data-baseweb="menu"] [role="option"]:hover,
[data-baseweb="menu"] li:hover {
    background: #eff6ff !important;
    color: #1d4ed8 !important;
}

/* Arrow icon */
[data-testid="stSelectbox"] svg {
    fill: #6b7280 !important;
    color: #6b7280 !important;
}

/* CHECKBOXES */
[data-testid="stCheckbox"] label,
[data-testid="stCheckbox"] span,
[data-testid="stCheckbox"] p {
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
    font-size: 0.95rem !important;
}

/* EXPANDERS */
[data-testid="stExpander"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    box-shadow: var(--shadow) !important;
}

[data-testid="stExpander"] summary,
[data-testid="stExpander"] summary span,
[data-testid="stExpander"] summary p {
    color: var(--text) !important;
    font-weight: 600 !important;
}

[data-testid="stExpander"] [data-testid="stExpanderDetails"] {
    background: #ffffff !important;
}

[data-testid="stExpander"] [data-testid="stExpanderDetails"] *:not(.notes-preview *) {
    color: var(--text) !important;
}

/* METRICS */
[data-testid="stMetric"] {
    background: #ffffff !important;
    border: 1px solid var(--border) !important;
    border-radius: 12px !important;
    padding: 1rem 1.2rem !important;
    box-shadow: var(--shadow) !important;
}

[data-testid="stMetricValue"],
[data-testid="stMetricValue"] > div {
    color: var(--accent) !important;
    font-family: 'DM Serif Display', serif !important;
    font-size: 1.8rem !important;
}

[data-testid="stMetricLabel"],
[data-testid="stMetricLabel"] > div,
[data-testid="stMetricLabel"] p {
    color: var(--muted) !important;
    font-size: 0.82rem !important;
    font-weight: 600 !important;
    text-transform: uppercase !important;
    letter-spacing: 0.06em !important;
}

/* ALERTS */
[data-testid="stAlert"] { border-radius: 10px !important; }
[data-testid="stAlert"] p,
[data-testid="stAlert"] span,
[data-testid="stAlert"] div { font-family: 'DM Sans', sans-serif !important; }

div[data-baseweb="notification"][kind="positive"] {
    background: #f0fdf4 !important;
    border-left: 4px solid #16a34a !important;
    color: #14532d !important;
}
div[data-baseweb="notification"][kind="warning"] {
    background: #fffbeb !important;
    border-left: 4px solid var(--warning) !important;
    color: #78350f !important;
}
div[data-baseweb="notification"][kind="negative"] {
    background: #fef2f2 !important;
    border-left: 4px solid #dc2626 !important;
    color: #7f1d1d !important;
}
div[data-baseweb="notification"][kind="info"] {
    background: #eff6ff !important;
    border-left: 4px solid var(--accent) !important;
    color: #1e3a8a !important;
}

/* PROGRESS BAR */
[data-testid="stProgressBar"] > div > div {
    background: linear-gradient(90deg, var(--accent), var(--accent2)) !important;
    border-radius: 99px !important;
}
[data-testid="stProgressBar"] > div {
    background: #dbeafe !important;
    border-radius: 99px !important;
}

/* DIVIDER */
hr {
    border: none !important;
    border-top: 1px solid var(--border) !important;
    margin: 1.2rem 0 !important;
}

/* MARKDOWN */
.stMarkdown p, .stMarkdown li, .stMarkdown span,
[data-testid="stMarkdownContainer"] p,
[data-testid="stMarkdownContainer"] li,
[data-testid="stMarkdownContainer"] span {
    color: var(--text) !important;
    font-family: 'DM Sans', sans-serif !important;
}
[data-testid="stMarkdownContainer"] strong {
    color: var(--text) !important;
    font-weight: 700 !important;
}
[data-testid="stMarkdownContainer"] code {
    background: #f1f5f9 !important;
    color: #1e3a8a !important;
    border-radius: 4px !important;
    padding: 0.15em 0.4em !important;
    font-size: 0.88em !important;
}

/* SPINNER */
[data-testid="stSpinner"] * { color: var(--accent) !important; }

/* CUSTOM COMPONENTS */
.hero-card {
    background: linear-gradient(135deg, #1e40af 0%, #3730a3 50%, #6d28d9 100%);
    border-radius: 18px;
    padding: 2rem 2.4rem;
    margin-bottom: 1.8rem;
    box-shadow: 0 8px 40px rgba(37,99,235,0.28);
    position: relative;
    overflow: hidden;
}
.hero-card::before {
    content: '';
    position: absolute; top: -40px; right: -40px;
    width: 180px; height: 180px;
    border-radius: 50%;
    background: rgba(255,255,255,0.07);
    pointer-events: none;
}
.hero-card::after {
    content: '';
    position: absolute; bottom: -60px; left: 30%;
    width: 240px; height: 240px;
    border-radius: 50%;
    background: rgba(255,255,255,0.04);
    pointer-events: none;
}
.hero-card h1 {
    color: #ffffff !important;
    font-size: 2.1rem !important;
    margin: 0 !important;
    font-family: 'DM Serif Display', serif !important;
    text-shadow: 0 1px 4px rgba(0,0,0,0.2);
}
.hero-card p {
    color: rgba(255,255,255,0.82) !important;
    margin: 0.5rem 0 0 0 !important;
    font-size: 1rem !important;
}

.metric-box {
    background: #ffffff;
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 1.4rem 1.2rem;
    text-align: center;
    box-shadow: var(--shadow);
    transition: transform 0.15s ease, box-shadow 0.15s ease;
}
.metric-box:hover { transform: translateY(-2px); box-shadow: var(--shadow-lg); }
.metric-box .metric-val {
    font-family: 'DM Serif Display', serif;
    font-size: 2.6rem;
    font-weight: 700;
    color: var(--accent);
    line-height: 1;
}
.metric-box .metric-label {
    font-size: 0.78rem;
    color: var(--muted);
    text-transform: uppercase;
    letter-spacing: 0.09em;
    margin-top: 0.4rem;
    font-weight: 600;
}

.section-header {
    font-family: 'DM Serif Display', serif !important;
    font-size: 1.2rem !important;
    font-weight: 400 !important;
    color: #111827 !important;
    margin: 1.6rem 0 0.9rem 0;
    padding-bottom: 0.5rem;
    border-bottom: 2px solid var(--border);
}

.topic-chip {
    display: inline-block;
    background: #eff6ff;
    border: 1px solid #93c5fd;
    color: #1d4ed8 !important;
    border-radius: 20px;
    padding: 0.25rem 0.9rem;
    font-size: 0.83rem;
    margin: 0.2rem;
    font-weight: 500;
}

.notes-preview {
    background: #ffffff;
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 1.8rem;
    white-space: pre-wrap;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.93rem;
    line-height: 1.8;
    max-height: 550px;
    overflow-y: auto;
    color: #1a1814 !important;
    box-shadow: var(--shadow);
}
.notes-preview::-webkit-scrollbar { width: 6px; }
.notes-preview::-webkit-scrollbar-track { background: #f0eeea; border-radius: 3px; }
.notes-preview::-webkit-scrollbar-thumb { background: #d1cdc8; border-radius: 3px; }

.plan-preview {
    background: #f8faff;
    border: 1px solid #bfdbfe;
    border-left: 4px solid var(--accent);
    border-radius: 14px;
    padding: 1.8rem;
    white-space: pre-wrap;
    font-family: 'DM Sans', sans-serif;
    font-size: 0.93rem;
    line-height: 1.8;
    max-height: 500px;
    overflow-y: auto;
    color: #1a1814 !important;
    box-shadow: var(--shadow);
}
.plan-preview::-webkit-scrollbar { width: 6px; }
.plan-preview::-webkit-scrollbar-track { background: #eff6ff; border-radius: 3px; }
.plan-preview::-webkit-scrollbar-thumb { background: #93c5fd; border-radius: 3px; }

.info-banner {
    background: #eff6ff;
    border: 1px solid #bfdbfe;
    border-left: 4px solid var(--accent);
    border-radius: 10px;
    padding: 0.9rem 1.2rem;
    font-size: 0.9rem;
    color: #1e40af !important;
    margin-bottom: 1rem;
}
.info-banner * { color: #1e40af !important; }

.success-banner {
    background: #f0fdf4;
    border: 1px solid #bbf7d0;
    border-left: 4px solid #16a34a;
    border-radius: 10px;
    padding: 0.9rem 1.2rem;
    font-size: 0.9rem;
    color: #14532d !important;
    margin-bottom: 1rem;
}
.success-banner * { color: #14532d !important; }

.warning-banner {
    background: #fffbeb;
    border: 1px solid #fde68a;
    border-left: 4px solid var(--warning);
    border-radius: 10px;
    padding: 0.9rem 1.2rem;
    font-size: 0.9rem;
    color: #78350f !important;
    margin-bottom: 1rem;
}
.warning-banner * { color: #78350f !important; }

.secure-badge {
    background: #f0fdf4;
    border: 1px solid #86efac;
    border-radius: 8px;
    padding: 0.6rem 1rem;
    font-size: 0.82rem;
    color: #15803d !important;
    margin-bottom: 1.2rem;
    font-weight: 500;
}
.secure-badge * { color: #15803d !important; }

.sidebar-logo {
    font-family: 'DM Serif Display', serif;
    font-size: 1.6rem;
    color: #2563eb !important;
    margin-bottom: 0.1rem;
    display: block;
}
.sidebar-sub {
    font-size: 0.78rem;
    color: #6b6560 !important;
    margin-bottom: 1rem;
    font-weight: 400;
    display: block;
}

.tag-pill {
    display: inline-block;
    padding: 0.2rem 0.75rem;
    border-radius: 20px;
    font-size: 0.78rem;
    font-weight: 600;
    letter-spacing: 0.03em;
}
.tag-blue   { background: #dbeafe; color: #1d4ed8 !important; }
.tag-green  { background: #dcfce7; color: #15803d !important; }
.tag-purple { background: #f3e8ff; color: #7e22ce !important; }
.tag-amber  { background: #fef3c7; color: #92400e !important; }

.example-topics {
    background: #f8faff;
    border: 1px solid #dbeafe;
    border-radius: 10px;
    padding: 1rem 1.2rem;
    margin-top: 0.5rem;
}
.example-topics p { color: #374151 !important; font-size: 0.85rem !important; margin: 0 !important; }

.level-badge {
    display: inline-block;
    padding: 0.3rem 0.9rem;
    border-radius: 20px;
    font-size: 0.82rem;
    font-weight: 700;
    letter-spacing: 0.04em;
}
.level-beginner  { background: #dcfce7; color: #15803d !important; }
.level-intermediate { background: #fef3c7; color: #92400e !important; }
.level-advanced  { background: #fce7f3; color: #9d174d !important; }

</style>
""", unsafe_allow_html=True)

# ──────────────────────────────────────────────
# API KEY
# ──────────────────────────────────────────────
api_key = st.secrets.get("GROQ_API_KEY", "")
if not api_key:
    st.error(
        "⚠️ **GROQ_API_KEY not found.**\n\n"
        "In Streamlit Cloud go to: **App Settings → Secrets** and add:\n\n"
        "```\nGROQ_API_KEY = \"gsk_your_key_here\"\n```"
    )
    st.stop()

# ──────────────────────────────────────────────
# SESSION STATE
# ──────────────────────────────────────────────
for key, default in {
    "notes": "",
    "study_plan": "",
    "topics_list": [],
    "subject_name": "",
    "level": "Beginner",
    "history": [],
    "current_topics_input": "",
}.items():
    if key not in st.session_state:
        st.session_state[key] = default

# ──────────────────────────────────────────────
# GROQ CLIENT
# ──────────────────────────────────────────────
@st.cache_resource
def get_groq_client(key: str):
    return Groq(api_key=key)


# ──────────────────────────────────────────────
# TEXT SANITIZER
# ──────────────────────────────────────────────
def sanitize_for_pdf(text: str) -> str:
    result = []
    for ch in text:
        code = ord(ch)
        if code < 256:
            result.append(ch)
        else:
            replacements = {
                '\u2018': "'", '\u2019': "'", '\u201c': '"', '\u201d': '"',
                '\u2013': '-', '\u2014': '--', '\u2026': '...', '\u2022': '*',
                '\u00b7': '*', '\u2192': '->', '\u2190': '<-',
                '\u25cf': '*', '\u25b6': '>',
            }
            result.append(replacements.get(ch, '?'))
    return ''.join(result)


def clean_markdown(text: str) -> str:
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    return text


# ──────────────────────────────────────────────
# LLM HELPERS
# ──────────────────────────────────────────────
def llm_call(client, system: str, user: str, max_tokens: int = 3000) -> str:
    response = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        messages=[
            {"role": "system", "content": system},
            {"role": "user",   "content": user},
        ],
        temperature=0.4,
        max_tokens=max_tokens,
    )
    return response.choices[0].message.content.strip()


def detect_input_type(client, user_input: str) -> dict:
    """Detect whether input is a full subject or specific topics."""
    system = (
        "You are an education expert. Given the user input, determine:\n"
        "1. Whether this is a FULL SUBJECT (like DSA, Python, Java, OOP, etc.) or SPECIFIC TOPICS (like queues, stacks, linked lists)\n"
        "2. Extract a clean subject/area name\n"
        "3. List all topics to cover\n\n"
        "Return ONLY valid JSON, no markdown, no explanation:\n"
        '{"type": "full_subject" or "specific_topics", "subject": "...", "topics": ["topic1", "topic2", ...]}\n\n'
        "For full subjects, list ALL major topics of that subject.\n"
        "For specific topics, list exactly what was provided.\n"
        "Use only ASCII characters."
    )
    raw = llm_call(client, system, f"User input: {user_input}", max_tokens=800)
    try:
        cleaned = raw.strip().replace("```json", "").replace("```", "").strip()
        start = cleaned.find("{")
        end   = cleaned.rfind("}")
        if start != -1 and end != -1:
            return json.loads(cleaned[start:end+1])
    except Exception:
        pass
    return {"type": "specific_topics", "subject": user_input, "topics": [user_input]}


def generate_notes(client, topics: list, subject: str, level: str) -> str:
    topics_str = ", ".join(topics)
    is_full = len(topics) > 5

    if is_full:
        system = (
            f"You are a world-class {subject} educator and textbook author. "
            f"Create comprehensive, complete lecture notes for the subject: {subject}. "
            f"Target audience: {level} level students.\n\n"
            "Structure the notes covering ALL these topics thoroughly:\n"
            f"{topics_str}\n\n"
            "For each topic include:\n"
            "- Clear explanation (what it is, why it matters)\n"
            "- Key concepts and definitions\n"
            "- How it works (step by step)\n"
            "- Real examples and use cases\n"
            "- Common mistakes to avoid\n\n"
            "Use this structure:\n"
            f"# Complete Notes: {subject}\n"
            "## Introduction\n"
            "## [Topic 1]\n"
            "### What is it?\n"
            "### Key Concepts\n"
            "### Examples\n"
            "## [Topic 2]\n"
            "... (continue for all topics)\n"
            "## Quick Reference Summary\n\n"
            "Be extremely thorough. Use markdown. ASCII only."
        )
    else:
        system = (
            f"You are a world-class educator and expert in {subject}. "
            f"Create comprehensive, detailed notes on the following topics: {topics_str}. "
            f"Target audience: {level} level students.\n\n"
            "For each topic provide:\n"
            "- Complete explanation from zero to pro\n"
            "- All key definitions and concepts\n"
            "- Detailed working/mechanism\n"
            "- Real-world examples and code (if applicable)\n"
            "- Visual representation using ASCII art/diagrams where helpful\n"
            "- Common interview questions about this topic\n"
            "- Key points to remember\n\n"
            f"# Detailed Notes: {', '.join(topics)}\n"
            "## [Topic Name]\n"
            "### Overview\n"
            "### Core Concepts\n"
            "### In-depth Explanation\n"
            "### Examples\n"
            "### Interview Questions\n"
            "### Key Takeaways\n\n"
            "Be extremely detailed and educational. ASCII only."
        )
    return llm_call(client, system, f"Generate complete notes for: {topics_str}", max_tokens=3000)


def generate_study_plan(client, topics: list, subject: str, level: str, days: int) -> str:
    topics_str = ", ".join(topics)
    system = (
        f"You are an expert academic planner and {subject} educator. "
        f"Create a detailed, realistic {days}-day study plan for a {level} level student "
        f"to master: {topics_str}.\n\n"
        "The plan should:\n"
        "- Be organized day by day\n"
        "- Include specific tasks and time estimates per day\n"
        "- Balance theory, practice, and revision\n"
        "- Include resource suggestions (types of resources, not specific URLs)\n"
        "- Include milestones and checkpoints\n"
        "- Include practice problem suggestions\n"
        "- Be achievable and progressive\n\n"
        "Format:\n"
        f"# {days}-Day Study Plan: {subject}\n"
        "## Overview & Goals\n"
        "## Daily Schedule\n"
        "### Day 1: [Focus Topic]\n"
        "- Morning (X hours): ...\n"
        "- Afternoon (X hours): ...\n"
        "- Evening (X hours): ...\n"
        "### Day 2: ...\n"
        "...\n"
        "## Milestones\n"
        "## Tips for Success\n\n"
        "ASCII only. Be specific and actionable."
    )
    return llm_call(client, system, f"Create study plan for: {topics_str}", max_tokens=2500)


# ──────────────────────────────────────────────
# PDF EXPORT
# ──────────────────────────────────────────────
class SmartNotesPDF(FPDF):
    def __init__(self, meta):
        super().__init__()
        self.meta = meta
        self.set_margins(20, 20, 20)
        self.set_auto_page_break(auto=True, margin=20)

    def header(self):
        self.set_fill_color(37, 99, 235)
        self.rect(0, 0, 210, 14, 'F')
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(255, 255, 255)
        self.set_y(3)
        self.cell(0, 8, "SmartNotes - AI Study Notes Generator", align="C")
        self.set_text_color(0, 0, 0)
        self.ln(10)

    def footer(self):
        self.set_y(-13)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 8, f"Page {self.page_no()} | Generated {self.meta.get('date', '')} | Level: {self.meta.get('level', '')}", align="C")


def export_pdf(content: str, meta: dict, doc_title: str = "Study Notes") -> bytes:
    pdf = SmartNotesPDF(meta)
    pdf.add_page()

    # A4 width=210mm, margins left=20, right=20 → usable=170mm
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin  # 170mm

    # ── Title block ──
    title_h = 26
    pdf.set_fill_color(239, 246, 255)
    pdf.set_draw_color(191, 219, 254)
    pdf.set_line_width(0.3)
    pdf.rect(pdf.l_margin, pdf.get_y(), usable_w, title_h, 'FD')
    pdf.set_font("Helvetica", "B", 14)
    pdf.set_text_color(30, 64, 175)
    pdf.set_y(pdf.get_y() + 4)
    safe_subject = sanitize_for_pdf(meta.get("subject") or doc_title)
    pdf.cell(usable_w, 9, safe_subject, align="C", ln=True)

    info_line = sanitize_for_pdf(
        f"Level: {meta.get('level', '-')}   |   Date: {meta.get('date', '-')}"
    )
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(100, 116, 139)
    pdf.cell(usable_w, 7, info_line, align="C", ln=True)
    pdf.ln(8)

    # Line height constants
    LH_BODY   = 6.5   # body text line height
    LH_H1     = 8.5
    LH_H2     = 7.5
    LH_H3     = 7.0
    LH_BULLET = 6.5

    def safe_multi_cell(w, h, txt):
        """Render multi_cell always starting at l_margin, reset x after."""
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(w, h, txt, border=0)

    def safe_bullet(indent_mm, bullet_char, txt, lh):
        """
        Render a bullet point safely:
        - Print bullet character at indent
        - Print text in a multi_cell that wraps correctly
        The trick: print bullet in a narrow cell, then go back up and
        render the text multi_cell offset by bullet width.
        """
        bullet_w  = 6.0
        text_x    = pdf.l_margin + indent_mm + bullet_w
        text_w    = usable_w - indent_mm - bullet_w

        y_start = pdf.get_y()

        # Bullet dot
        pdf.set_x(pdf.l_margin + indent_mm)
        pdf.cell(bullet_w, lh, bullet_char, border=0)

        # Now cursor is at (text_x, y_start) — but cell() moved x
        # Move back to text_x on same line
        pdf.set_xy(text_x, y_start)
        pdf.multi_cell(text_w, lh, txt, border=0)

    for raw_line in content.split("\n"):
        line = sanitize_for_pdf(raw_line)
        stripped = line.strip()

        if not stripped:
            pdf.ln(2)
            continue

        # ── H1 ──
        if stripped.startswith("# "):
            ct = clean_markdown(stripped[2:]).strip()
            if not ct:
                continue
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 13)
            pdf.set_text_color(30, 64, 175)
            safe_multi_cell(usable_w, LH_H1, ct)
            # Underline
            pdf.set_draw_color(147, 197, 253)
            pdf.set_line_width(0.5)
            y = pdf.get_y()
            pdf.line(pdf.l_margin, y, pdf.l_margin + usable_w, y)
            pdf.ln(3)

        # ── H2 ──
        elif stripped.startswith("## "):
            ct = clean_markdown(stripped[3:]).strip()
            if not ct:
                continue
            pdf.ln(2)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(37, 99, 235)
            safe_multi_cell(usable_w, LH_H2, ct)
            pdf.ln(1)

        # ── H3 ──
        elif stripped.startswith("### "):
            ct = clean_markdown(stripped[4:]).strip()
            if not ct:
                continue
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(55, 65, 81)
            safe_multi_cell(usable_w, LH_H3, ct)

        # ── H4 ──
        elif stripped.startswith("#### "):
            ct = clean_markdown(stripped[5:]).strip()
            if not ct:
                continue
            pdf.set_font("Helvetica", "BI", 10)
            pdf.set_text_color(75, 85, 99)
            safe_multi_cell(usable_w, LH_H3, ct)

        # ── Sub-bullet (3+ spaces or tab before - or *) ──
        elif (
            line.startswith("   - ") or line.startswith("   * ")
            or line.startswith("\t- ") or line.startswith("\t* ")
        ):
            ct = clean_markdown(stripped.lstrip("-* ")).strip()
            if not ct:
                continue
            pdf.set_font("Helvetica", "", 9)
            pdf.set_text_color(80, 80, 80)
            safe_bullet(10, "-", ct, LH_BULLET)

        # ── Top-level bullet ──
        elif stripped.startswith("- ") or stripped.startswith("* "):
            ct = clean_markdown(stripped[2:]).strip()
            if not ct:
                continue
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            safe_bullet(3, chr(149), ct, LH_BULLET)

        # ── Numbered list ──
        elif re.match(r'^\d+[\.\)]\s', stripped):
            ct = clean_markdown(re.sub(r'^\d+[\.\)]\s+', '', stripped)).strip()
            num = re.match(r'^(\d+[\.\)])', stripped).group(1)
            if not ct:
                continue
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            safe_bullet(3, num, ct, LH_BULLET)

        # ── Body text ──
        else:
            ct = clean_markdown(stripped)
            if not ct:
                continue
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            safe_multi_cell(usable_w, LH_BODY, ct)

    return bytes(pdf.output())


# ──────────────────────────────────────────────
# DOCX EXPORT
# ──────────────────────────────────────────────
def export_docx(content: str, meta: dict, doc_title: str = "Study Notes") -> bytes:
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    title = doc.add_heading(f"SmartNotes — {doc_title}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title.runs:
        title.runs[0].font.color.rgb = RGBColor(30, 64, 175)
        title.runs[0].font.size = Pt(22)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info = (
        "Subject: " + (meta.get("subject") or "-") +
        "   |   Level: " + (meta.get("level") or "-") +
        "   |   Date: " + (meta.get("date") or "-")
    )
    run = p.add_run(info)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.font.size = Pt(10)
    doc.add_paragraph()

    for line in content.split("\n"):
        ls = line.strip()
        if not ls:
            doc.add_paragraph()
            continue
        if ls.startswith("# "):
            h = doc.add_heading(ls[2:], level=1)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(30, 64, 175)
                h.runs[0].font.size = Pt(16)
        elif ls.startswith("## "):
            h = doc.add_heading(ls[3:], level=2)
            if h.runs:
                h.runs[0].font.color.rgb = RGBColor(37, 99, 235)
                h.runs[0].font.size = Pt(13)
        elif ls.startswith("### "):
            h = doc.add_heading(ls[4:], level=3)
            if h.runs:
                h.runs[0].font.size = Pt(12)
        elif ls.startswith("- ") or ls.startswith("* "):
            p = doc.add_paragraph(ls[2:], style="List Bullet")
            if p.runs:
                p.runs[0].font.size = Pt(11)
        else:
            p = doc.add_paragraph(ls)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ──────────────────────────────────────────────
# SIDEBAR
# ──────────────────────────────────────────────
with st.sidebar:
    st.markdown('<div class="sidebar-logo">📚 SmartNotes</div>', unsafe_allow_html=True)
    st.markdown('<div class="sidebar-sub">AI-Powered Study Notes Generator</div>', unsafe_allow_html=True)

    st.markdown(
        '<div class="secure-badge">🔒 API key loaded securely from Secrets</div>',
        unsafe_allow_html=True,
    )

    st.markdown("#### ⚙️ Generation Settings")

    level = st.selectbox(
        "📊 Knowledge Level",
        ["Beginner", "Intermediate", "Advanced"],
        index=0,
        key="level_select",
    )
    st.session_state.level = level

    study_days = st.selectbox(
        "📅 Study Plan Duration",
        [3, 5, 7, 10, 14, 21, 30],
        index=1,
        key="study_days",
    )

    st.divider()

    st.markdown("#### 🎛️ What to Generate")
    gen_notes = st.checkbox("📝 Generate Notes", value=True)
    gen_plan  = st.checkbox("📅 Generate Study Plan", value=True)

    st.divider()

    if st.session_state.history:
        st.markdown("#### 📊 Quick Stats")
        total_sessions = len(st.session_state.history)
        total_topics   = sum(s["topic_count"] for s in st.session_state.history)

        cols = st.columns(2)
        cols[0].metric("Sessions", total_sessions)
        cols[1].metric("Topics",   total_topics)
        st.divider()

    st.markdown(
        "<p style='font-size:0.76rem;color:#94a3b8;margin-top:0.5rem'>"
        "Powered by Groq · LLaMA 3.1 · AI Notes Engine"
        "</p>",
        unsafe_allow_html=True,
    )


# ──────────────────────────────────────────────
# HERO SECTION
# ──────────────────────────────────────────────
st.markdown("""
<div class="hero-card">
  <h1>📚 AI Study Notes &amp; Plan Generator</h1>
  <p>Enter any subject or topics → get comprehensive notes from zero to pro + a personalized study plan + PDF &amp; DOCX export</p>
</div>
""", unsafe_allow_html=True)

tab1, tab2, tab3, tab4 = st.tabs([
    "✏️  Enter Topics",
    "📝  Study Notes",
    "📅  Study Plan",
    "📚  History",
])

# ══════════════════════════════════════════════
# TAB 1 — Enter Topics
# ══════════════════════════════════════════════
with tab1:
    col_left, col_right = st.columns([1.15, 0.85], gap="large")

    with col_left:
        st.markdown('<div class="section-header">✏️ Enter Your Subject or Topics</div>', unsafe_allow_html=True)

        topics_input = st.text_area(
            "Topics / Subject Input",
            height=160,
            placeholder=(
                "Examples:\n"
                "• DSA  (generates full Data Structures & Algorithms notes)\n"
                "• Python  (generates complete Python notes)\n"
                "• OOP, Inheritance, Polymorphism, Encapsulation\n"
                "• Stack, Queue, Linked List\n"
                "• Java Collections Framework"
            ),
            label_visibility="collapsed",
            key="topics_input_area",
        )

        st.markdown(
            '<div class="example-topics">'
            '<p>💡 <strong>Full subjects:</strong> DSA, Python, Java, OOP, DBMS, OS, Computer Networks, Machine Learning</p>'
            '<p style="margin-top:0.4rem">💡 <strong>Specific topics:</strong> Binary Search Trees, Recursion, Dynamic Programming, Sorting Algorithms</p>'
            '</div>',
            unsafe_allow_html=True,
        )

        st.markdown('<div style="height:0.8rem"></div>', unsafe_allow_html=True)
        generate_btn = st.button("🚀 Generate Notes & Plan", use_container_width=True)

    with col_right:
        st.markdown('<div class="section-header">ℹ️ How It Works</div>', unsafe_allow_html=True)

        st.markdown("""
<div class="info-banner">
<strong>Step 1:</strong> Type a full subject (e.g. <em>DSA</em>) or specific topics (e.g. <em>Stack, Queue</em>)<br><br>
<strong>Step 2:</strong> Choose your level (Beginner / Intermediate / Advanced) in the sidebar<br><br>
<strong>Step 3:</strong> Click Generate — the AI creates complete notes & a study plan<br><br>
<strong>Step 4:</strong> Download as PDF or DOCX from the Notes & Plan tabs
</div>
""", unsafe_allow_html=True)

        st.markdown('<div class="section-header">📊 Current Settings</div>', unsafe_allow_html=True)

        level_class = {
            "Beginner": "level-beginner",
            "Intermediate": "level-intermediate",
            "Advanced": "level-advanced",
        }.get(st.session_state.level, "level-beginner")

        st.markdown(
            f'<div style="margin-bottom:0.7rem">'
            f'<span class="level-badge {level_class}">📊 {st.session_state.level}</span>'
            f'<span class="tag-pill tag-blue" style="margin-left:0.5rem">📅 {study_days}-day plan</span>'
            f'</div>',
            unsafe_allow_html=True,
        )

        if st.session_state.notes:
            topic_count = len(st.session_state.topics_list)
            notes_wc    = len(st.session_state.notes.split())
            plan_wc     = len(st.session_state.study_plan.split()) if st.session_state.study_plan else 0

            st.markdown('<div class="section-header">📊 Last Generated</div>', unsafe_allow_html=True)
            mc1, mc2 = st.columns(2)
            with mc1:
                st.markdown(f'<div class="metric-box"><div class="metric-val">{topic_count}</div><div class="metric-label">Topics</div></div>', unsafe_allow_html=True)
            with mc2:
                st.markdown(f'<div class="metric-box"><div class="metric-val">{notes_wc:,}</div><div class="metric-label">Words</div></div>', unsafe_allow_html=True)

    # ── Generate logic ──
    if generate_btn:
        raw_input = topics_input.strip()
        if not raw_input:
            st.warning("⚠️ Please enter at least one subject or topic before generating.")
        else:
            client = get_groq_client(api_key)
            progress_bar = st.progress(0, text="🧠 Analyzing your input...")

            try:
                # Step 1: Detect type and extract topics
                progress_bar.progress(15, text="🔍 Detecting subject and topics...")
                detected = detect_input_type(client, raw_input)
                subject  = detected.get("subject", raw_input)
                topics   = detected.get("topics", [raw_input])
                if not topics:
                    topics = [raw_input]

                st.session_state.subject_name = subject
                st.session_state.topics_list  = topics
                st.session_state.current_topics_input = raw_input

                # Step 2: Generate notes
                if gen_notes:
                    progress_bar.progress(35, text="📝 Generating comprehensive notes...")
                    st.session_state.notes = generate_notes(
                        client, topics, subject, st.session_state.level
                    )
                else:
                    st.session_state.notes = ""

                # Step 3: Generate study plan
                if gen_plan:
                    progress_bar.progress(70, text="📅 Building your study plan...")
                    st.session_state.study_plan = generate_study_plan(
                        client, topics, subject, st.session_state.level, study_days
                    )
                else:
                    st.session_state.study_plan = ""

                progress_bar.progress(100, text="✅ Done!")

                # Save to history
                st.session_state.history.append({
                    "date":        datetime.today().strftime("%Y-%m-%d"),
                    "subject":     subject,
                    "topic_count": len(topics),
                    "level":       st.session_state.level,
                    "has_notes":   bool(st.session_state.notes),
                    "has_plan":    bool(st.session_state.study_plan),
                    "topics":      topics[:8],
                })

                st.balloons()
                st.success(
                    f"✅ Generated notes for **{subject}** covering **{len(topics)} topics**! "
                    "Switch to the **Study Notes** and **Study Plan** tabs to view and download."
                )

            except Exception as e:
                st.error(f"❌ Generation failed: {e}")

    # Show detected topics if available
    if st.session_state.topics_list:
        st.markdown('<div class="section-header">🏷️ Detected Topics</div>', unsafe_allow_html=True)
        chips = "".join(f'<span class="topic-chip">📌 {t}</span>' for t in st.session_state.topics_list)
        st.markdown(f'<div style="line-height:2.4">{chips}</div>', unsafe_allow_html=True)


# ══════════════════════════════════════════════
# TAB 2 — Study Notes
# ══════════════════════════════════════════════
with tab2:
    if not st.session_state.notes:
        st.markdown(
            '<div class="info-banner">ℹ️ Enter a subject or topics in the <strong>Enter Topics</strong> tab and click Generate to create notes here.</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="section-header">📝 Generated Study Notes</div>', unsafe_allow_html=True)

        meta = {
            "subject": st.session_state.subject_name,
            "level":   st.session_state.level,
            "date":    datetime.today().strftime("%Y-%m-%d"),
        }

        exp_col1, exp_col2, exp_col3 = st.columns([1, 1, 2])
        with exp_col1:
            try:
                pdf_bytes = export_pdf(st.session_state.notes, meta, f"Notes - {st.session_state.subject_name}")
                st.download_button(
                    "📄 Download PDF",
                    data=pdf_bytes,
                    file_name=f"notes_{st.session_state.subject_name.replace(' ', '_')}_{meta['date']}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF export error: {e}")

        with exp_col2:
            try:
                docx_bytes = export_docx(st.session_state.notes, meta, f"Notes - {st.session_state.subject_name}")
                st.download_button(
                    "📝 Download DOCX",
                    data=docx_bytes,
                    file_name=f"notes_{st.session_state.subject_name.replace(' ', '_')}_{meta['date']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"DOCX export error: {e}")

        with exp_col3:
            level_class = {
                "Beginner": "level-beginner",
                "Intermediate": "level-intermediate",
                "Advanced": "level-advanced",
            }.get(st.session_state.level, "level-beginner")
            st.markdown(
                f'<div style="padding:0.5rem 0">'
                f'<span class="level-badge {level_class}">📊 {st.session_state.level}</span> '
                f'<span class="tag-pill tag-purple" style="margin-left:0.4rem">🎓 {st.session_state.subject_name}</span> '
                f'<span class="tag-pill tag-blue" style="margin-left:0.4rem">📌 {len(st.session_state.topics_list)} topics</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

        st.markdown('<div style="height:0.6rem"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="notes-preview">' +
            st.session_state.notes.replace("\n", "<br>") +
            '</div>',
            unsafe_allow_html=True,
        )


# ══════════════════════════════════════════════
# TAB 3 — Study Plan
# ══════════════════════════════════════════════
with tab3:
    if not st.session_state.study_plan:
        st.markdown(
            '<div class="info-banner">ℹ️ Enter a subject or topics in the <strong>Enter Topics</strong> tab and click Generate to create a study plan here.</div>',
            unsafe_allow_html=True,
        )
    else:
        st.markdown('<div class="section-header">📅 Your Personalized Study Plan</div>', unsafe_allow_html=True)

        meta = {
            "subject": st.session_state.subject_name,
            "level":   st.session_state.level,
            "date":    datetime.today().strftime("%Y-%m-%d"),
        }

        plan_col1, plan_col2, plan_col3 = st.columns([1, 1, 2])
        with plan_col1:
            try:
                plan_pdf = export_pdf(
                    st.session_state.study_plan, meta,
                    f"Study Plan - {st.session_state.subject_name}"
                )
                st.download_button(
                    "📄 Download Plan PDF",
                    data=plan_pdf,
                    file_name=f"study_plan_{st.session_state.subject_name.replace(' ', '_')}_{meta['date']}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"PDF export error: {e}")

        with plan_col2:
            try:
                plan_docx = export_docx(
                    st.session_state.study_plan, meta,
                    f"Study Plan - {st.session_state.subject_name}"
                )
                st.download_button(
                    "📝 Download Plan DOCX",
                    data=plan_docx,
                    file_name=f"study_plan_{st.session_state.subject_name.replace(' ', '_')}_{meta['date']}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"DOCX export error: {e}")

        with plan_col3:
            level_class = {
                "Beginner": "level-beginner",
                "Intermediate": "level-intermediate",
                "Advanced": "level-advanced",
            }.get(st.session_state.level, "level-beginner")
            st.markdown(
                f'<div style="padding:0.5rem 0">'
                f'<span class="level-badge {level_class}">📊 {st.session_state.level}</span> '
                f'<span class="tag-pill tag-green" style="margin-left:0.4rem">🎯 {st.session_state.subject_name}</span>'
                f'</div>',
                unsafe_allow_html=True,
            )

        st.markdown('<div style="height:0.6rem"></div>', unsafe_allow_html=True)
        st.markdown(
            '<div class="plan-preview">' +
            st.session_state.study_plan.replace("\n", "<br>") +
            '</div>',
            unsafe_allow_html=True,
        )


# ══════════════════════════════════════════════
# TAB 4 — History
# ══════════════════════════════════════════════
with tab4:
    st.markdown('<div class="section-header">📚 Generation History</div>', unsafe_allow_html=True)

    if not st.session_state.history:
        st.markdown(
            '<div class="info-banner">ℹ️ No notes generated yet. Start by entering a subject in the <strong>Enter Topics</strong> tab.</div>',
            unsafe_allow_html=True,
        )
    else:
        h_col1, h_col2, h_col3, h_col4 = st.columns(4)
        with h_col1:
            st.markdown(f'<div class="metric-box"><div class="metric-val">{len(st.session_state.history)}</div><div class="metric-label">Sessions</div></div>', unsafe_allow_html=True)
        with h_col2:
            total_t = sum(s["topic_count"] for s in st.session_state.history)
            st.markdown(f'<div class="metric-box"><div class="metric-val">{total_t}</div><div class="metric-label">Topics Covered</div></div>', unsafe_allow_html=True)
        with h_col3:
            with_notes = sum(1 for s in st.session_state.history if s["has_notes"])
            st.markdown(f'<div class="metric-box"><div class="metric-val">{with_notes}</div><div class="metric-label">Notes Generated</div></div>', unsafe_allow_html=True)
        with h_col4:
            with_plans = sum(1 for s in st.session_state.history if s["has_plan"])
            st.markdown(f'<div class="metric-box"><div class="metric-val">{with_plans}</div><div class="metric-label">Plans Generated</div></div>', unsafe_allow_html=True)

        st.markdown('<div style="height:1rem"></div>', unsafe_allow_html=True)

        for i, session in enumerate(reversed(st.session_state.history)):
            session_num = len(st.session_state.history) - i
            notes_tag = '<span class="tag-pill tag-green">✓ Notes</span>' if session["has_notes"] else '<span class="tag-pill tag-amber">No Notes</span>'
            plan_tag  = '<span class="tag-pill tag-blue">✓ Plan</span>'  if session["has_plan"]  else '<span class="tag-pill tag-amber">No Plan</span>'

            level_class = {
                "Beginner": "level-beginner",
                "Intermediate": "level-intermediate",
                "Advanced": "level-advanced",
            }.get(session.get("level", "Beginner"), "level-beginner")

            with st.expander(
                f"#{session_num}  |  {session['date']}  —  {session['subject']}  ({session['level']})",
                expanded=(i == 0),
            ):
                ec1, ec2, ec3, ec4 = st.columns(4)
                ec1.metric("Date",    session["date"])
                ec2.metric("Subject", session["subject"])
                ec3.metric("Topics",  session["topic_count"])
                ec4.metric("Level",   session["level"])

                if session.get("topics"):
                    chips = "".join(f'<span class="topic-chip">📌 {t}</span>' for t in session["topics"])
                    st.markdown(f'<div style="margin-top:0.5rem;line-height:2.4">{chips}</div>', unsafe_allow_html=True)

                st.markdown(
                    f'<div style="margin-top:0.6rem">{notes_tag} {plan_tag} '
                    f'<span class="level-badge {level_class}" style="margin-left:0.4rem">{session["level"]}</span>'
                    f'<span class="tag-pill tag-purple" style="margin-left:0.4rem">Session #{session_num}</span></div>',
                    unsafe_allow_html=True,
                )
